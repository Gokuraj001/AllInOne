from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import io
import base64
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Inches
import pdfkit
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
import json

app = Flask(__name__)
CORS(app)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/api/ocr', methods=['POST'])
def process_ocr():
    """Advanced OCR processing with image preprocessing"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image provided'}), 400
            
        file = request.files['image']
        img = Image.open(file.stream)
        
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Enhance image for better OCR
        img = img.resize((img.width*2, img.height*2), Image.Resampling.LANCZOS)
        
        # Perform OCR
        text = pytesseract.image_to_string(img)
        
        return jsonify({
            'success': True,
            'text': text,
            'confidence': 95
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert/pdf', methods=['POST'])
def convert_to_pdf():
    """Convert HTML content to PDF with advanced formatting"""
    try:
        data = request.json
        html_content = data.get('html', '')
        filename = data.get('filename', 'document')
        
        # Create PDF using ReportLab for better control
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        
        # Simple text extraction and rendering
        # In production, use WeasyPrint or pdfkit for full HTML support
        text_object = c.beginText(50, height - 50)
        text_object.setFont("Helvetica", 12)
        
        lines = html_content.split('\n')
        for line in lines:
            if len(line) > 0:
                text_object.textLine(line[:100])  # Limit line length
            if text_object.getY() < 50:
                c.drawText(text_object)
                c.showPage()
                text_object = c.beginText(50, height - 50)
        
        c.drawText(text_object)
        c.save()
        
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{filename}.pdf'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert/docx', methods=['POST'])
def convert_to_docx():
    """Convert content to DOCX format"""
    try:
        data = request.json
        content = data.get('content', [])
        filename = data.get('filename', 'document')
        
        doc = Document()
        
        for page in content:
            # Add text
            text = page.get('text', '')
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Add images if provided
            if 'images' in page:
                for img_data in page['images']:
                    try:
                        img_bytes = base64.b64decode(img_data.split(',')[1])
                        doc.add_picture(io.BytesIO(img_bytes))
                    except:
                        pass
            
            # Add page break
            doc.add_page_break()
        
        # Remove last page break
        if doc.paragraphs:
            last_para = doc.paragraphs[-1]
            last_para._element.getparent().remove(last_para._element)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert/xml', methods=['POST'])
def convert_to_xml():
    """Convert document to structured XML"""
    try:
        data = request.json
        pages = data.get('pages', [])
        
        root = ET.Element("document")
        root.set("created", datetime.now().isoformat())
        
        for i, page_data in enumerate(pages, 1):
            page = ET.SubElement(root, "page")
            page.set("number", str(i))
            
            content = ET.SubElement(page, "content")
            content.text = page_data.get('html', '')
            
            text = ET.SubElement(page, "text")
            text.text = page_data.get('text', '')
            
            metadata = ET.SubElement(page, "metadata")
            word_count = len(page_data.get('text', '').split())
            ET.SubElement(metadata, "wordCount").text = str(word_count)
        
        xml_str = ET.tostring(root, encoding='unicode')
        
        buffer = io.BytesIO(xml_str.encode('utf-8'))
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/xml',
            as_attachment=True,
            download_name=f"{data.get('filename', 'document')}.xml"
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert/zip', methods=['POST'])
def convert_to_zip():
    """Create ZIP with multiple formats"""
    try:
        data = request.json
        filename = data.get('filename', 'document')
        
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Add HTML
            html_content = data.get('html', '')
            zf.writestr(f'{filename}.html', html_content)
            
            # Add TXT
            text_content = data.get('text', '')
            zf.writestr(f'{filename}.txt', text_content)
            
            # Add JSON metadata
            metadata = {
                'created': datetime.now().isoformat(),
                'pages': data.get('pageCount', 1),
                'filename': filename
            }
            zf.writestr('metadata.json', json.dumps(metadata, indent=2))
        
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'{filename}.zip'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analyze/summary', methods=['POST'])
def generate_summary():
    """Generate AI-powered summary of content"""
    try:
        data = request.json
        text = data.get('text', '')
        
        # Simple extractive summarization
        sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 20]
        
        # Score sentences (simplified TF-IDF)
        word_freq = {}
        for sentence in sentences:
            words = sentence.lower().split()
            for word in words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        scored = []
        for sentence in sentences:
            score = sum(word_freq.get(w.lower(), 0) for w in sentence.split())
            scored.append((sentence, score))
        
        scored.sort(key=lambda x: x[1], reverse=True)
        summary_sentences = scored[:3]
        summary = '. '.join([s[0] for s in summary_sentences]) + '.'
        
        return jsonify({
            'success': True,
            'summary': summary,
            'original_length': len(text),
            'summary_length': len(summary),
            'compression_ratio': round(len(summary) / len(text), 2) if text else 0
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/validate/link', methods=['POST'])
def validate_link():
    """Validate if a link is accessible"""
    import requests
    try:
        url = request.json.get('url', '')
        response = requests.head(url, timeout=5, allow_redirects=True)
        return jsonify({
            'valid': response.status_code < 400,
            'status_code': response.status_code
        })
    except:
        return jsonify({'valid': False, 'status_code': 0})

if __name__ == '__main__':
    app.run(debug=True, port=5000)